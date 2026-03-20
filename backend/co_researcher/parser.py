import os
import time
import requests

LLAMAPARSE_API_KEY = os.getenv("LLAMAPARSE_API_KEY", "llx-kxyquEwhrd9z5QeQWtrGeHh2dwbAzqDz34nP1dSh4qo6iAhL")
LLAMAPARSE_BASE_URL = "https://api.cloud.llamaindex.ai/api/v1/parsing"

def parse_pdf(file_bytes: bytes, filename: str) -> str:
    """
    Parse a PDF using LlamaParse API.
    Returns extracted text as markdown.
    """
    headers = {
        "Authorization": f"Bearer {LLAMAPARSE_API_KEY}",
    }

    # Step 1: Upload the file
    files = {
        "file": (filename, file_bytes, "application/pdf"),
    }
    data = {
        "result_type": "markdown",
        "language": "en",
    }

    upload_resp = requests.post(
        f"{LLAMAPARSE_BASE_URL}/upload",
        headers=headers,
        files=files,
        data=data,
        timeout=60
    )
    upload_resp.raise_for_status()
    job_id = upload_resp.json()["id"]

    # Step 2: Poll for completion (max 5 minutes)
    for _ in range(60):
        status_resp = requests.get(
            f"{LLAMAPARSE_BASE_URL}/job/{job_id}",
            headers=headers,
            timeout=30
        )
        status_resp.raise_for_status()
        status = status_resp.json()["status"]

        if status == "SUCCESS":
            break
        elif status == "ERROR":
            raise RuntimeError(f"LlamaParse failed for {filename}: {status_resp.json()}")

        time.sleep(5)
    else:
        raise TimeoutError(f"LlamaParse timed out for {filename}")

    # Step 3: Download result
    result_resp = requests.get(
        f"{LLAMAPARSE_BASE_URL}/job/{job_id}/result/markdown",
        headers=headers,
        timeout=30
    )
    result_resp.raise_for_status()
    return result_resp.json()["markdown"]


def parse_docx(file_bytes: bytes, filename: str) -> str:
    """Parse DOCX using python-docx. Returns text content as markdown."""
    from docx import Document
    import io
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name.startswith('Heading'):
            level_str = para.style.name.replace('Heading ', '').strip()
            try:
                level = int(level_str)
            except ValueError:
                level = 2
            paragraphs.append(f"{'#' * level} {text}")
        else:
            paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
        if rows:
            # Add header separator after first row
            header_sep = '| ' + ' | '.join(['---'] * len(table.rows[0].cells)) + ' |'
            rows.insert(1, header_sep)
            paragraphs.append('\n'.join(rows))

    return '\n\n'.join(paragraphs)


def parse_document(file_bytes: bytes, filename: str) -> str:
    """Parse PDF, DOCX, or plain text. Returns extracted text as markdown."""
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext == 'pdf':
        return parse_pdf(file_bytes, filename)
    elif ext in ('docx', 'doc'):
        return parse_docx(file_bytes, filename)
    elif ext == 'txt':
        # Plain text (e.g., research description) - return as-is
        return file_bytes.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Upload PDF or DOCX.")
