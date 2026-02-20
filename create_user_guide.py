#!/usr/bin/env python3
"""
Create comprehensive user guide for 2nd Brain
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_user_guide():
    doc = Document()

    # Set document title
    title = doc.add_heading('2nd Brain', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete User Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Version info
    version = doc.add_paragraph('Version 2.0 | February 2026')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. What is 2nd Brain?',
        '2. Getting Started',
        '3. The Dashboard',
        '4. Connecting Your Data Sources',
        '5. Managing Your Documents',
        '6. Using the AI Chat Assistant',
        '7. Knowledge Gaps',
        '8. Training Guides & Videos',
        '9. Settings & Preferences',
        '10. Troubleshooting',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ===== SECTION 1: WHAT IS 2ND BRAIN =====
    doc.add_heading('1. What is 2nd Brain?', level=1)

    doc.add_paragraph(
        '2nd Brain is an AI-powered knowledge management system designed to help organizations '
        'preserve, organize, and access their collective knowledge. Think of it as a smart assistant '
        'that remembers everything your team knows and can answer questions about it.'
    )

    doc.add_heading('The Problem It Solves', level=2)
    doc.add_paragraph(
        'When employees leave a company or when information is scattered across different tools '
        '(emails, Slack, documents, etc.), valuable knowledge gets lost. 2nd Brain solves this by:'
    )

    bullets = [
        'Automatically collecting information from all your tools',
        'Organizing and classifying documents using AI',
        'Making everything searchable with natural language questions',
        'Identifying gaps in your knowledge base',
        'Helping train new team members faster',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('Key Features Overview', level=2)

    features_table = doc.add_table(rows=1, cols=2)
    features_table.style = 'Table Grid'
    header_cells = features_table.rows[0].cells
    header_cells[0].text = 'Feature'
    header_cells[1].text = 'What It Does'

    features = [
        ('AI Chat Assistant', 'Ask questions in plain English and get answers from your knowledge base'),
        ('Smart Integrations', 'Connect Gmail, Slack, Google Drive, OneDrive, Notion, GitHub, and more'),
        ('Document Management', 'Automatically organize and classify your documents'),
        ('Knowledge Gap Detection', 'Find missing information in your documentation'),
        ('Training Guide Generator', 'Create training materials automatically'),
        ('Email Notifications', 'Get notified when syncs complete'),
    ]

    for feature, description in features:
        row = features_table.add_row()
        row.cells[0].text = feature
        row.cells[1].text = description

    doc.add_page_break()

    # ===== SECTION 2: GETTING STARTED =====
    doc.add_heading('2. Getting Started', level=1)

    doc.add_heading('Creating Your Account', level=2)
    steps = [
        ('Step 1', 'Go to the 2nd Brain website and click "Sign Up"'),
        ('Step 2', 'Enter your email address and create a password'),
        ('Step 3', 'Enter your organization name (this creates your workspace)'),
        ('Step 4', 'Check your email for a verification link and click it'),
        ('Step 5', 'Log in to access your dashboard'),
    ]
    for step, desc in steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(desc)

    doc.add_heading('First-Time Setup', level=2)
    doc.add_paragraph(
        'After logging in for the first time, we recommend following these steps:'
    )

    setup_steps = [
        'Connect at least one data source (like Gmail or Google Drive)',
        'Wait for the initial sync to complete',
        'Review your documents in the Documents page',
        'Try asking a question in the Chat interface',
    ]
    for i, step in enumerate(setup_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ===== SECTION 3: THE DASHBOARD =====
    doc.add_heading('3. The Dashboard', level=1)

    doc.add_paragraph(
        'The dashboard is your home base in 2nd Brain. Here\'s what you\'ll find:'
    )

    doc.add_heading('Sidebar Navigation', level=2)
    nav_items = [
        ('ChatBot', 'Your AI assistant - ask questions about your knowledge base'),
        ('Documents', 'View, search, and manage all your imported documents'),
        ('Knowledge Gaps', 'See what information might be missing from your docs'),
        ('Integrations', 'Connect and manage your data sources'),
        ('Training Guides', 'Create and view training materials'),
        ('Analytics', 'View usage statistics and insights'),
        ('Settings', 'Manage your account and preferences'),
    ]

    for nav, desc in nav_items:
        p = doc.add_paragraph()
        p.add_run(f'{nav}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ===== SECTION 4: CONNECTING DATA SOURCES =====
    doc.add_heading('4. Connecting Your Data Sources', level=1)

    doc.add_paragraph(
        '2nd Brain can pull information from many different sources. Each connection is called an '
        '"integration." Here\'s how to set them up:'
    )

    doc.add_heading('Available Integrations', level=2)

    integrations_table = doc.add_table(rows=1, cols=3)
    integrations_table.style = 'Table Grid'
    header = integrations_table.rows[0].cells
    header[0].text = 'Integration'
    header[1].text = 'What It Imports'
    header[2].text = 'How to Connect'

    integrations = [
        ('Gmail', 'Your emails and attachments', 'Click Connect → Sign in with Google'),
        ('Google Drive', 'Documents, spreadsheets, presentations', 'Click Connect → Sign in with Google'),
        ('Slack', 'Channel messages and shared files', 'Click Connect → Authorize in Slack'),
        ('Microsoft OneDrive', 'Word, Excel, PowerPoint files', 'Click Connect → Sign in with Microsoft'),
        ('Notion', 'Pages and databases', 'Click Connect → Authorize in Notion'),
        ('GitHub', 'Code repositories and README files', 'Click Connect → Authorize with GitHub'),
        ('Box', 'Files and folders', 'Click Connect → Sign in with Box'),
        ('Zotero', 'Research papers and citations', 'Enter your API key'),
        ('Outlook', 'Emails from Microsoft account', 'Click Connect → Sign in with Microsoft'),
        ('Website Scraper', 'Any website content', 'Enter the website URL'),
        ('Email Forwarding', 'Forward emails to a special address', 'Copy your unique forwarding address'),
    ]

    for name, imports, connect in integrations:
        row = integrations_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = imports
        row.cells[2].text = connect

    doc.add_heading('How to Connect an Integration', level=2)

    connect_steps = [
        'Go to the Integrations page from the sidebar',
        'Find the integration you want to connect',
        'Click the "Connect" button',
        'Follow the prompts to authorize access (usually signing into that service)',
        'Once connected, click "Sync" to start importing your data',
    ]
    for i, step in enumerate(connect_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Sync Progress & Email Notifications', level=2)
    doc.add_paragraph(
        'When you sync an integration, a progress window appears showing:'
    )
    bullets = [
        'How many items were found',
        'How many have been processed',
        'Current status (crawling, extracting, embedding)',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Email Notification: ').bold = True
    p.add_run(
        'Check the "Email me when sync completes" box to receive an email when all your syncs '
        'are finished. This is useful for large imports that take a while.'
    )

    doc.add_page_break()

    # ===== SECTION 5: MANAGING DOCUMENTS =====
    doc.add_heading('5. Managing Your Documents', level=1)

    doc.add_paragraph(
        'The Documents page shows all the content that has been imported into your knowledge base.'
    )

    doc.add_heading('Document List', level=2)
    doc.add_paragraph('Each document shows:')
    bullets = [
        'Title - The name of the document or email subject',
        'Source - Where it came from (Gmail, Slack, Drive, etc.)',
        'Date - When it was created or received',
        'Classification - Work or Personal (assigned by AI)',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('Document Classification', level=2)
    doc.add_paragraph(
        '2nd Brain uses AI to automatically classify documents as either "Work" (relevant to your '
        'organization) or "Personal" (not work-related). You can:'
    )
    bullets = [
        'Confirm a classification if it\'s correct',
        'Reject a document if it was misclassified',
        'Bulk confirm multiple documents at once',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('Searching Documents', level=2)
    doc.add_paragraph(
        'Use the search bar to find specific documents. You can search by:'
    )
    bullets = [
        'Title or content keywords',
        'Sender name or email',
        'Date range',
        'Source type (Gmail, Slack, etc.)',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ===== SECTION 6: AI CHAT ASSISTANT =====
    doc.add_heading('6. Using the AI Chat Assistant', level=1)

    doc.add_paragraph(
        'The Chat feature is the heart of 2nd Brain. It lets you ask questions in plain English '
        'and get answers based on your organization\'s knowledge.'
    )

    doc.add_heading('How It Works', level=2)
    steps = [
        'Type your question in the chat box (just like texting)',
        'The AI searches through all your documents to find relevant information',
        'You get an answer with citations showing where the information came from',
        'Click on citations to see the original source documents',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Example Questions You Can Ask', level=2)
    examples = [
        '"What is our company\'s vacation policy?"',
        '"How do I submit an expense report?"',
        '"What were the key decisions from last month\'s board meeting?"',
        '"Who is responsible for the marketing budget?"',
        '"What is the process for onboarding new employees?"',
        '"Summarize the project status updates from this week"',
    ]
    for ex in examples:
        doc.add_paragraph(ex, style='List Bullet')

    doc.add_heading('Tips for Better Results', level=2)
    tips = [
        'Be specific in your questions',
        'Include relevant context (project names, dates, people)',
        'If you don\'t get a good answer, try rephrasing your question',
        'Use the feedback buttons (thumbs up/down) to help improve results',
    ]
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_heading('Voice Input', level=2)
    doc.add_paragraph(
        'You can also ask questions using your voice! Click the microphone icon and speak your '
        'question. The AI will transcribe what you said and search for an answer.'
    )

    doc.add_page_break()

    # ===== SECTION 7: KNOWLEDGE GAPS =====
    doc.add_heading('7. Knowledge Gaps', level=1)

    doc.add_paragraph(
        'Knowledge Gaps are areas where your documentation might be incomplete or missing important '
        'information. 2nd Brain automatically detects these gaps by analyzing your documents.'
    )

    doc.add_heading('What Are Knowledge Gaps?', level=2)
    doc.add_paragraph('The system looks for things like:')
    gaps = [
        'Decisions mentioned without explaining why they were made',
        'Processes referenced but not documented step-by-step',
        'People mentioned as the only source of certain knowledge ("Ask John, he knows")',
        'Vague timelines ("we\'ll do this eventually")',
        'Undefined terms or acronyms',
        'Contradictions between different documents',
    ]
    for gap in gaps:
        doc.add_paragraph(gap, style='List Bullet')

    doc.add_heading('How to Use Knowledge Gaps', level=2)
    steps = [
        'Go to the Knowledge Gaps page',
        'Click "Find Gaps" to analyze your documents',
        'Review the list of detected gaps',
        'For each gap, you can:',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    sub_items = [
        'Answer the question to fill in the missing knowledge',
        'Record a voice answer if typing is inconvenient',
        'Mark it as "not applicable" if it\'s not relevant',
        'Provide feedback on whether the gap was useful',
    ]
    for item in sub_items:
        doc.add_paragraph(f'   • {item}')

    doc.add_heading('Gap Categories', level=2)
    categories = [
        ('Missing Rationale', 'A decision was made but the "why" is not documented'),
        ('Missing Owner', 'A process exists but nobody is assigned to maintain it'),
        ('Bus Factor Risk', 'Critical knowledge is locked in one person\'s head'),
        ('Undefined Process', 'Something is done "the usual way" but not written down'),
        ('Contradiction', 'Different documents say different things'),
    ]
    for cat, desc in categories:
        p = doc.add_paragraph()
        p.add_run(f'{cat}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ===== SECTION 8: TRAINING GUIDES =====
    doc.add_heading('8. Training Guides & Videos', level=1)

    doc.add_paragraph(
        'Training Guides help you create educational materials for your team based on your '
        'knowledge base.'
    )

    doc.add_heading('Creating a Training Guide', level=2)
    steps = [
        'Go to the Training Guides page',
        'Click "Create New Guide"',
        'Enter a topic (e.g., "How to use our CRM system")',
        'The AI will generate a guide using information from your documents',
        'Review and edit the guide as needed',
        'Share with your team',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Video Generation', level=2)
    doc.add_paragraph(
        '2nd Brain can also create training videos with AI-generated narration:'
    )
    bullets = [
        'Videos are created from your training guide content',
        'AI voice reads the content aloud',
        'Slides are automatically generated',
        'Download videos to share with your team',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ===== SECTION 9: SETTINGS =====
    doc.add_heading('9. Settings & Preferences', level=1)

    doc.add_heading('Account Settings', level=2)
    settings = [
        ('Profile', 'Update your name and profile picture'),
        ('Email', 'Change your email address'),
        ('Password', 'Update your password'),
        ('Notifications', 'Control which emails you receive'),
    ]
    for setting, desc in settings:
        p = doc.add_paragraph()
        p.add_run(f'{setting}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Organization Settings (Admin Only)', level=2)
    admin_settings = [
        ('Team Members', 'Invite new users or remove existing ones'),
        ('Integrations', 'Manage which data sources are connected'),
        ('Data Retention', 'Control how long data is kept'),
    ]
    for setting, desc in admin_settings:
        p = doc.add_paragraph()
        p.add_run(f'{setting}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ===== SECTION 10: TROUBLESHOOTING =====
    doc.add_heading('10. Troubleshooting', level=1)

    doc.add_heading('Common Issues', level=2)

    issues = [
        ('Sync is taking a long time',
         'Large imports (hundreds of files) can take 10-30 minutes. Check the "Email me when done" '
         'box and you\'ll be notified when it\'s complete.'),
        ('I can\'t connect an integration',
         'Make sure you\'re signed into the correct account. Try disconnecting and reconnecting. '
         'Check that you\'ve authorized all required permissions.'),
        ('The AI chat isn\'t finding information I know exists',
         'The document might not have synced yet. Check the Documents page to confirm it\'s there. '
         'Try using different keywords in your question.'),
        ('I\'m getting too many irrelevant results',
         'Use the classification feature to mark personal documents as "Personal" so they\'re '
         'excluded from work searches.'),
        ('Knowledge Gaps seem wrong or unhelpful',
         'Use the feedback buttons to mark gaps as "not useful" - this helps improve future detection.'),
    ]

    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(f'Problem: {issue}').bold = True
        doc.add_paragraph(f'Solution: {solution}')
        doc.add_paragraph()

    doc.add_heading('Getting Help', level=2)
    doc.add_paragraph('If you need additional help:')
    bullets = [
        'Check our FAQ at help.use2ndbrain.com',
        'Email support at support@use2ndbrain.com',
        'Use the in-app feedback button to report issues',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # ===== FOOTER =====
    doc.add_page_break()
    doc.add_heading('Quick Reference Card', level=1)

    ref_table = doc.add_table(rows=1, cols=2)
    ref_table.style = 'Table Grid'
    header = ref_table.rows[0].cells
    header[0].text = 'Task'
    header[1].text = 'How To Do It'

    quick_ref = [
        ('Ask a question', 'Go to ChatBot → Type your question → Press Enter'),
        ('Connect Gmail', 'Integrations → Gmail → Connect → Sign in with Google'),
        ('Sync data', 'Integrations → Choose source → Click Sync'),
        ('Find documents', 'Documents → Use search bar or filters'),
        ('Find knowledge gaps', 'Knowledge Gaps → Click "Find Gaps"'),
        ('Create training guide', 'Training Guides → Create New → Enter topic'),
        ('Get email notifications', 'During sync → Check "Email me when done"'),
        ('Change settings', 'Click Settings in sidebar'),
    ]

    for task, how in quick_ref:
        row = ref_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = how

    # Save the document
    output_path = '/Users/badri/2ndBRAINPRANAV/2nd_Brain_User_Guide.docx'
    doc.save(output_path)
    print(f'✅ User guide created: {output_path}')
    return output_path

if __name__ == '__main__':
    create_user_guide()
